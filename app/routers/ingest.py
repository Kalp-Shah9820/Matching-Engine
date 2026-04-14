import uuid
import io
import json
import re

import openpyxl
import docx
import pdfplumber
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session

from app.database import get_db
from app.schemas import JobDescription, Candidate
from app.models import (
    JobDescriptionIn,
    CandidateIn,
    BulkCandidateIn,
    BulkJDIn,
    JobDescriptionOut,
    CandidateOut,
    UploadResponse,
)

router = APIRouter(prefix="/ingest", tags=["Ingestion"])


# ─── helper ────────────────────────────────────────────────────────────────────

def _safe_bool(value) -> bool | None:
    """Convert truthy Excel cell values to Python bool safely."""
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in ("true", "yes", "1")
    return bool(value)


def _safe_float(value) -> float | None:
    try:
        return float(value) if value is not None else None
    except (ValueError, TypeError):
        return None


def _extract_text_from_docx(contents: bytes) -> str:
    """Extracts text from DOCX while preserving line breaks and basic structure."""
    document = docx.Document(io.BytesIO(contents))
    # Preserve paragraph structure and handle table text if any
    full_text = []
    for para in document.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    # Also extract from tables (often JDs use tables)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)
                
    return "\n".join(full_text)


def _extract_text_from_pdf(contents: bytes) -> str:
    """Extracts text from PDF using pdfplumber with improved layout preservation."""
    text = []
    try:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages:
                # Use layout=True to maintain columns and alignments
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                text.append(page_text)
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        return ""
    return "\n".join(text)


def _extract_section(text: str, keywords: list[str]) -> str | None:
    """
    Smarter section extraction. 
    Finds a keyword and captures text until the next major heading or multiple line breaks.
    """
    lines = text.splitlines()
    lower_keywords = [k.lower() for k in keywords]
    
    start_idx = -1
    found_line_num = -1
    
    for i, line in enumerate(lines):
        clean_line = line.strip().lower()
        if any(k in clean_line for k in lower_keywords):
            # Check if this line is a heading (usually short)
            if len(clean_line) < 50:
                start_idx = i
                found_line_num = i
                break
                
    if start_idx == -1: return None
    
    section_content = []
    # Start collecting from the line AFTER the heading or the same line if it has content
    first_line_content = lines[start_idx].strip()
    # If the heading line has a colon, take what's after it
    if ':' in first_line_content:
        after_colon = first_line_content.split(':', 1)[1].strip()
        if after_colon:
            section_content.append(after_colon)
    
    # Look for the end of the section
    for i in range(start_idx + 1, len(lines)):
        line = lines[i].strip()
        if not line:
            # Allow up to 2 empty lines before assuming section change
            if i + 1 < len(lines) and lines[i+1].strip():
                continue
            else:
                if len(section_content) > 3: # If we have some content, a break might mean end
                    break
                continue
        
        # Heuristic: If we hit another line that looks like a heading, stop
        # A heading is usually short, all caps, or ends in a colon
        if (line.isupper() and len(line) < 40) or (line.endswith(':') and len(line) < 30):
            break
            
        section_content.append(line)
        
        # Stop if we hit common 'next' section keywords
        lower_line = line.lower()
        if any(stop in lower_line for stop in ['responsibilities', 'qualifications', 'requirements', 'about us']):
            if len(line) < 30: # Only if it looks like a header
                section_content.pop() # Remove the next header
                break
                
    return "\n".join(section_content).strip() or None


def _parse_job_text(text: str) -> tuple[str, str, str | None, str | None]:
    """Robustly parses JD text into structured components."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    
    # 1. Improved Title Detection
    title = 'Job Description'
    if lines:
        for i in range(min(5, len(lines))):
            line = lines[i]
            # Match explicit title labels
            title_match = re.search(
                r'(?:Job\s+Title|Position|Role|Opening|Vacancy)\s*[:\-]\s*(.+)',
                line, re.IGNORECASE
            )
            if title_match:
                title = title_match.group(1).strip()
                break
            # Fallback: Treat first line as title if it's short and no label found
            if i == 0 and len(line) < 60 and not line.endswith('.'):
                title = line
                break

    # 2. Extract Sections with prioritized keywords
    overview = _extract_section(text, ['overview', 'job summary', 'about the role', 'company overview', 'introduction'])
    if not overview and len(lines) > 2:
        # Fallback for overview: take first few non-title lines
        overview = "\n".join(lines[1:4])

    skills = _extract_section(text, ['required skills', 'technical skills', 'skills set', 'must-have', 'key skills', 'competencies'])
    
    # Requirements usually include things like responsibilities and qualifications
    requirements = _extract_section(text, ['requirements', 'qualifications', 'what you will do', 'responsibilities', 'expectations'])

    return title, overview or "", requirements, skills


def _parse_multiple_jobs(text: str) -> list[tuple[str, str, str | None, str | None]]:
    """
    Parse text that may contain multiple job descriptions.
    Splits only on explicit job-heading patterns to avoid false positives.
    """
    normalized_text = text.replace('\r\n', '\n').replace('\r', '\n')
    title_marker = re.compile(
        r'(?m)^(?:\d+\.|\*+\s*)?\s*(?:Job\s+Title|Position|Role|Opening|Vacancy|Job\s+Description|JD)\s*[:\-]',
        re.IGNORECASE,
    )

    matches = list(title_marker.finditer(normalized_text))
    if len(matches) < 2:
        # if the file only has one job description or no explicit headings, parse as a single JD
        return [_parse_job_text(normalized_text)]

    job_starts = [m.start() for m in matches]
    job_sections = []
    for idx, start in enumerate(job_starts):
        end = job_starts[idx + 1] if idx + 1 < len(job_starts) else len(normalized_text)
        section = normalized_text[start:end].strip()
        if section:
            job_sections.append(section)

    return [_parse_job_text(section) for section in job_sections] if job_sections else [_parse_job_text(normalized_text)]


def _row_to_candidate(row: dict) -> Candidate:
    """
    Maps a raw dict (from Excel row or JSON payload) → Candidate ORM object.
    Handles missing/null fields gracefully so partial profiles don't crash.
    """
    candidate_id = str(row.get("id") or row.get("candidate_id") or uuid.uuid4())

    return Candidate(
        candidate_id          = candidate_id,
        submission_id         = str(row.get("submission_id") or ""),
        name                  = row.get("name"),
        location              = row.get("location"),
        current_title         = row.get("current_title"),
        years_of_experience   = _safe_float(row.get("years_of_experience")
                                            or row.get("parsed_metadata_calculated_years_experience")),
        education_status      = row.get("education_status"),
        looking_for           = str(row.get("looking_for") or ""),
        current_position      = row.get("current_position"),
        current_company       = row.get("current_company"),
        degree                = row.get("degree"),
        notice_period         = row.get("notice_period"),
        open_to_working_at    = row.get("open_to_working_at"),
        gen_ai_experience     = row.get("gen_ai_experience"),
        expected_salary       = _safe_float(row.get("expected_salary")),
        engineer_type         = row.get("engineer_type"),
        programming_languages = row.get("programming_languages"),
        backend_frameworks    = row.get("backend_frameworks"),
        frontend_technologies = row.get("frontend_technologies"),
        mobile_technologies   = row.get("mobile_technologies"),
        parsed_summary        = row.get("parsed_summary"),
        parsed_skills         = row.get("parsed_skills"),
        parsed_work_experience= row.get("parsed_work_experience"),
        parsed_education      = str(row.get("parsed_metadata_education")
                                    or row.get("parsed_education") or ""),
        is_actively_looking   = _safe_bool(row.get("is_actively_or_passively_looking")
                                           or row.get("is_actively_looking")),
        currently_employed    = _safe_bool(row.get("currently_employed")),
        recent_company_type   = str(row.get("recent_experience_type") or ""),
        recent_company_funded = _safe_bool(row.get("most_recent_company_is_funded")),
        recent_company_size   = row.get("most_recent_company_size"),
    )


# ─── JOB DESCRIPTION endpoints ─────────────────────────────────────────────────

@router.post(
    "/job",
    response_model=JobDescriptionOut,
    status_code=status.HTTP_201_CREATED,
    summary="Add a single job description",
)
def add_job(payload: JobDescriptionIn, db: Session = Depends(get_db)):
    """
    Add one job description via JSON body.
    Returns the saved JD with its auto-generated jd_id.
    """
    jd = JobDescription(
        jd_id             = str(uuid.uuid4()),
        title             = payload.title,
        company           = payload.company,
        overview          = payload.overview,
        core_requirements = payload.core_requirements,
        preferred_quals   = payload.preferred_quals,
        responsibilities  = payload.responsibilities,
        required_skills   = payload.required_skills,
        employment_type   = payload.employment_type,
        location          = payload.location,
    )
    db.add(jd)
    db.commit()
    db.refresh(jd)
    return jd


@router.post(
    "/jobs/upload-file",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Upload a Word or PDF job description file",
)
async def upload_job_description_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    if not file.filename.endswith((".docx", ".pdf")):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx or .pdf files are accepted for job description upload.",
        )
    contents = await file.read()
    if file.filename.lower().endswith('.docx'):
        text = _extract_text_from_docx(contents)
    else:
        text = _extract_text_from_pdf(contents)

    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file did not contain any readable job description text.",
        )

    # Parse multiple jobs from the text
    parsed_jobs = _parse_multiple_jobs(text)
    
    saved_ids = []
    for title, overview, core_requirements, required_skills in parsed_jobs:
        jd = JobDescription(
            jd_id             = str(uuid.uuid4()),
            title             = title,
            company           = None,
            overview          = overview,
            core_requirements = core_requirements,
            preferred_quals   = None,
            responsibilities  = None,
            required_skills   = required_skills,
            employment_type   = None,
            location          = None,
        )
        db.add(jd)
        saved_ids.append(jd.jd_id)

    db.commit()
    return UploadResponse(
        message=f"{len(saved_ids)} job description(s) ingested from {file.filename}.",
        count=len(saved_ids),
        ids=saved_ids,
    )


@router.post(
    "/jobs/bulk",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Add multiple job descriptions at once",
)
def add_jobs_bulk(payload: BulkJDIn, db: Session = Depends(get_db)):
    """
    Add multiple JDs in one request.
    Useful for seeding the system with all available roles.
    """
    saved_ids = []
    for item in payload.job_descriptions:
        jd = JobDescription(
            jd_id             = str(uuid.uuid4()),
            title             = item.title,
            company           = item.company,
            overview          = item.overview,
            core_requirements = item.core_requirements,
            preferred_quals   = item.preferred_quals,
            responsibilities  = item.responsibilities,
            required_skills   = item.required_skills,
            employment_type   = item.employment_type,
            location          = item.location,
        )
        db.add(jd)
        saved_ids.append(jd.jd_id)

    db.commit()
    return UploadResponse(
        message=f"{len(saved_ids)} job description(s) saved successfully.",
        count=len(saved_ids),
        ids=saved_ids,
    )


@router.get(
    "/jobs",
    response_model=list[JobDescriptionOut],
    summary="List all job descriptions",
)
def list_jobs(db: Session = Depends(get_db)):
    return db.query(JobDescription).all()


@router.delete(
    "/jobs/{jd_id}",
    status_code=status.HTTP_200_OK,
    summary="Delete a job description by ID",
)
def delete_job(jd_id: str, db: Session = Depends(get_db)):
    """
    Delete a job description by its jd_id.
    Returns 404 if the job doesn't exist.
    """
    jd = db.query(JobDescription).filter(JobDescription.jd_id == jd_id).first()
    if not jd:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Job description with id {jd_id} not found",
        )
    db.delete(jd)
    db.commit()
    return {"message": f"Job description {jd_id} deleted."}



# ─── CANDIDATE endpoints ────────────────────────────────────────────────────────

@router.post(
    "/candidate",
    response_model=CandidateOut,
    status_code=status.HTTP_201_CREATED,
    summary="Add a single candidate",
)
def add_candidate(payload: CandidateIn, db: Session = Depends(get_db)):
    """
    Add one candidate via JSON body.
    All fields except candidate_id are optional to handle incomplete profiles.
    """
    row = payload.model_dump()
    candidate = _row_to_candidate(row)
    db.add(candidate)
    db.commit()
    db.refresh(candidate)
    return candidate


@router.post(
    "/candidates/bulk",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Add multiple candidates via JSON",
)
def add_candidates_bulk(payload: BulkCandidateIn, db: Session = Depends(get_db)):
    """
    Add a list of candidates in one request.
    Skips duplicates (same candidate_id) with a warning rather than crashing.
    """
    saved_ids = []
    skipped = 0

    for item in payload.candidates:
        row = item.model_dump()
        candidate = _row_to_candidate(row)

        existing = db.query(Candidate).filter(
            Candidate.candidate_id == candidate.candidate_id
        ).first()

        if existing:
            skipped += 1
            continue

        db.add(candidate)
        saved_ids.append(candidate.candidate_id)

    db.commit()

    msg = f"{len(saved_ids)} candidate(s) saved."
    if skipped:
        msg += f" {skipped} duplicate(s) skipped."

    return UploadResponse(message=msg, count=len(saved_ids), ids=saved_ids)


@router.post(
    "/candidates/upload-excel",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Bulk upload candidates from Excel (.xlsx)",
)
async def upload_candidates_excel(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """
    Upload the Sample_Candidate_Data.xlsx directly.
    - Reads all rows from the first sheet.
    - Maps columns to Candidate fields automatically.
    - Skips rows with duplicate candidate IDs.
    - Skips completely empty rows silently.
    """
    # Validate file type
    if not file.filename.endswith((".xlsx", ".xlsm")):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .xlsx or .xlsm files are accepted.",
        )

    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=True)
    ws = wb.active

    rows = ws.iter_rows(values_only=True)
    headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(next(rows))]

    saved_ids = []
    skipped_dupes = 0
    skipped_empty = 0

    for raw_row in rows:
        # Skip completely empty rows
        if all(v is None for v in raw_row):
            skipped_empty += 1
            continue

        row = dict(zip(headers, raw_row))
        candidate = _row_to_candidate(row)

        existing = db.query(Candidate).filter(
            Candidate.candidate_id == candidate.candidate_id
        ).first()

        if existing:
            skipped_dupes += 1
            continue

        db.add(candidate)
        saved_ids.append(candidate.candidate_id)

    db.commit()

    msg = f"{len(saved_ids)} candidate(s) imported from Excel."
    if skipped_dupes:
        msg += f" {skipped_dupes} duplicate(s) skipped."
    if skipped_empty:
        msg += f" {skipped_empty} empty row(s) ignored."

    return UploadResponse(message=msg, count=len(saved_ids), ids=saved_ids)


@router.get(
    "/candidates",
    response_model=list[CandidateOut],
    summary="List all candidates (paginated)",
)
def list_candidates(skip: int = 0, limit: int = 50, db: Session = Depends(get_db)):
    """
    Returns candidates with pagination.
    Default: first 50. Use ?skip=50&limit=50 for the next page.
    """
    return db.query(Candidate).offset(skip).limit(limit).all()